from docx import Document

# Crear un documento nuevo
doc = Document()
doc.add_heading("Reporte automático", level=1)
doc.add_paragraph("Este documento fue generado con Python.")
doc.save("reporte.docx")

# Leer un documento existente
doc = Document("archivo.docx")
for p in doc.paragraphs:
    print(p.text)